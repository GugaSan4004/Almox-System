from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH

class init:
    
    def __init__(self, path):
        self.path = path
        
    def generate_return(self, data: dict):
        doc = Document(self.path + r"\static\py\doc\12_dez.docx")

        table = doc.tables[0]

        rows = table.rows

        data_items = list(data.items())

        for index, (ar_code, values) in enumerate(data_items, start=1):

            if index >= len(rows):
                break

            motivo = values[0].upper()
            destinatario = values[1].upper()

            cells = rows[index].cells
            
            cells[0].text = destinatario
            cells[1].text = ar_code
            cells[2].text = motivo
            
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if index == len(data_items) and index + 1 < len(rows):
                next_cells = rows[index + 1].cells

                next_cells[0].text = "***"
                next_cells[1].text = "***"
                next_cells[2].text = "***"
                
                for cell in next_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        for p in doc.paragraphs:
            if "DOCUMENTO EMITIDO" in p.text:
                p.clear()
                r1 = p.add_run("                                                    DOCUMENTO EMITIDO\n")
                r2 = p.add_run(f"                                                EM: {datetime.now().strftime("%d/%m/%Y")}")
                
                r1.bold = True
                r2.bold = True
                
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER           
                break

        doc.save(self.path + r"\static\py\doc\saved.docx")
        return self.path + r"\static\py\doc\saved.docx"